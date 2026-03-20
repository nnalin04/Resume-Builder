"""
ATS-safe DOCX resume generator.

Rules:
  - Single column, no tables, no text boxes, no columns
  - Calibri 11pt body, 20pt name
  - Proper paragraph styles so ATS parsers see clean heading/body hierarchy
  - Returns raw bytes so the FastAPI endpoint can stream directly
"""
import io
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Colour palette ────────────────────────────────────────────────────────────

_HEADING_COLOUR = RGBColor(0x1A, 0x56, 0xDB)   # brand blue


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_font(run, size_pt: float, bold: bool = False, italic: bool = False, colour: RGBColor | None = None):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour


def _section_heading(doc: Document, title: str):
    """Bold blue heading with a thin bottom border — looks clean, ATS-reads as plain text."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)

    run = para.add_run(title.upper())
    _set_font(run, 10.5, bold=True, colour=_HEADING_COLOUR)

    # Thin bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A56DB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _body_para(doc: Document, text: str = "", bold: bool = False, italic: bool = False,
               size_pt: float = 11, space_before: float = 0, space_after: float = 2,
               align=WD_ALIGN_PARAGRAPH.LEFT) -> Any:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.alignment = align
    if text:
        run = para.add_run(text)
        _set_font(run, size_pt, bold=bold, italic=italic)
    return para


def _bullet(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    _set_font(run, 11)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_resume_docx(sections: dict) -> bytes:
    """Convert resume sections dict to ATS-safe DOCX bytes."""
    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Global default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    contact = sections.get("contact") or {}

    # ── Name ──────────────────────────────────────────────────────────────────
    name = contact.get("name", "").strip()
    if name:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(name)
        _set_font(run, 20, bold=True)

    # ── Contact line ─────────────────────────────────────────────────────────
    contact_fields = ["email", "phone", "location", "linkedin", "github", "website"]
    contact_parts = [contact.get(f, "").strip() for f in contact_fields if contact.get(f, "").strip()]
    if contact_parts:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run("  |  ".join(contact_parts))
        _set_font(run, 9.5)

    # ── Summary ───────────────────────────────────────────────────────────────
    summary = (sections.get("summary") or "").strip()
    if summary:
        _section_heading(doc, "Summary")
        _body_para(doc, summary, space_after=4)

    # ── Experience ────────────────────────────────────────────────────────────
    experience = sections.get("experience") or []
    if experience:
        _section_heading(doc, "Experience")
        for job in experience:
            company = (job.get("company") or "").strip()
            title = (job.get("title") or "").strip()
            location = (job.get("location") or "").strip()
            start = (job.get("start_date") or "").strip()
            end = (job.get("end_date") or "Present").strip()
            date_str = f"{start} – {end}".strip(" –")

            # Company + date
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(0)
            r1 = para.add_run(company)
            _set_font(r1, 11, bold=True)
            if date_str:
                r2 = para.add_run(f"    {date_str}")
                _set_font(r2, 10.5)

            # Title + location
            title_loc = ", ".join(p for p in [title, location] if p)
            if title_loc:
                _body_para(doc, title_loc, italic=True, space_before=0, space_after=2)

            # Bullets
            for bullet in (job.get("bullets") or []):
                _bullet(doc, bullet)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = sections.get("skills") or {}
    skill_lines = []
    if isinstance(skills, dict):
        for cat, items in skills.items():
            if items:
                vals = ", ".join(items) if isinstance(items, list) else str(items)
                skill_lines.append(f"{cat.capitalize()}: {vals}")
    elif isinstance(skills, list):
        skill_lines.append(", ".join(str(s) for s in skills))

    if skill_lines:
        _section_heading(doc, "Skills")
        for line in skill_lines:
            _body_para(doc, line, space_before=1, space_after=1)

    # ── Education ─────────────────────────────────────────────────────────────
    education = sections.get("education") or []
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            institution = (edu.get("institution") or "").strip()
            grad_date = (edu.get("graduation_date") or "").strip()
            degree = (edu.get("degree") or "").strip()
            field = (edu.get("field") or "").strip()
            gpa = (edu.get("gpa") or "")

            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after = Pt(0)
            r1 = para.add_run(institution)
            _set_font(r1, 11, bold=True)
            if grad_date:
                r2 = para.add_run(f"    {grad_date}")
                _set_font(r2, 10.5)

            degree_parts = [p for p in [degree, field] if p]
            if gpa:
                degree_parts.append(f"GPA: {gpa}")
            if degree_parts:
                _body_para(doc, ", ".join(degree_parts), italic=True, space_before=0, space_after=3)

    # ── Projects ──────────────────────────────────────────────────────────────
    projects = sections.get("projects") or []
    if projects:
        _section_heading(doc, "Projects")
        for proj in projects:
            proj_name = (proj.get("name") or "").strip()
            tech_stack = proj.get("tech_stack") or []
            tech = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
            description = (proj.get("description") or "").strip()
            link = (proj.get("link") or "").strip()

            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after = Pt(0)
            r1 = para.add_run(proj_name)
            _set_font(r1, 11, bold=True)
            if tech:
                r2 = para.add_run(f"  ·  {tech}")
                _set_font(r2, 10.5, italic=True)

            if description:
                _body_para(doc, description, space_before=1, space_after=1)
            if link:
                _body_para(doc, link, size_pt=9.5, space_before=0, space_after=3)

    # ── Certifications ────────────────────────────────────────────────────────
    certifications = sections.get("certifications") or []
    if certifications:
        _section_heading(doc, "Certifications")
        for cert in certifications:
            parts = [
                (cert.get("name") or "").strip(),
                (cert.get("issuer") or "").strip(),
                (cert.get("date") or "").strip(),
            ]
            _body_para(doc, "  |  ".join(p for p in parts if p), space_before=2, space_after=2)

    # ── Serialize ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
